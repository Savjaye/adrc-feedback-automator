import pandas as pd

import docx

import os
from datetime import datetime as dt

# open interpretations and conclusions. place name. #MAKE SURE YOU REMEMBER TO ENCRYPT
fdbk_files = sorted(os.listdir("./input/feedback_template"))
fdbk_files = [file for file in fdbk_files if file != ".DS_Store"]
fdbk_files = [file for file in fdbk_files if file != "Archive"]
#   get a list of all the ids with a template... if they don't have a template, we need to make a severe letter
fdbk_ids = [int(fdbk_file[:4]) for fdbk_file in fdbk_files if not fdbk_file.startswith('~$08')] 
print(fdbk_ids)

# save all data to a dictionary
participant_data = pd.read_csv("./input/feedback_data.csv")
participant_data_dic = participant_data.to_dict()

participant_data["visit_date"] = pd.to_datetime(participant_data["visit_date"])

participant_data = participant_data.assign(visit_year=participant_data["visit_date"].dt.year)

adrcid = participant_data["Regtryid"].tolist()
subject_salutations = participant_data["subject_salutations"].tolist()
visit_date = participant_data["visit_date"].tolist()
visit_year = participant_data["visit_year"].tolist()
visit_number = participant_data["visit_number"].tolist()
language = participant_data["testing_language"].tolist()
street_address = participant_data["street_address"].tolist()
city = participant_data["city_state_zip"].tolist()

# open interpretations and conclusions. place name.
int_and_con_files = sorted(os.listdir("./input/interpretations_and_conclusions"))
int_and_con_files = [file for file in int_and_con_files if file != "Archive"]
int_and_con_files = [file for file in int_and_con_files if file != ".DS_Store"]
# int_and_con_ids = [id[:4] for id in int_and_con_files]
print(int_and_con_files)
# print(int_and_con_ids)
counter = 0
for file in int_and_con_files:
    interpretations_and_conclusions = docx.Document(f"./input/interpretations_and_conclusions/{file}")
    for paragraph in interpretations_and_conclusions.paragraphs:
        if "Name:" in paragraph.text:
            paragraph.text = paragraph.text.replace("Name:", f"Name: {subject_salutations[counter]}")
    interpretations_and_conclusions.save(
        f"./output/interpretations_and_conclusions_final/{adrcid[counter]}-{visit_year[counter]}.docx")
    counter += 1

templates = []
# make new letter 
    #   filter through subjects in df
for participant in participant_data_dic['Regtryid']:
    templates = []
    #    grab the appropriate template
    if adrcid[participant] in fdbk_ids:
        templates.append(docx.Document(f"./templates/Feedback_temp_english.docx"))
        if language[participant].lower() == 'spanish':
            templates.append(docx.Document(f"./templates/Feedback_temp_spanish.docx")) 
    else:
        templates.append(docx.Document(f"./templates/Feedback_temp_severe.docx"))
        print(f"!!!! {adrcid[participant]} is severe")

    print(street_address[participant])
    for i, template in enumerate(templates): 
        for paragraph in template.paragraphs:
            if "[name]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[name]", subject_salutations[participant])
            elif "[street_address]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[street_address]",
                                                        street_address[participant])
            elif "[city]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[city]", city[participant])
            elif "[date]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[date]", visit_date[participant].strftime("%m/%d/%Y"))
            elif "[visit]" in paragraph.text:
                paragraph.text = paragraph.text.replace("[visit]", str(visit_number[participant]))
        if i == 0:
            template.save(f"./Output/feedback_final/{int(adrcid[participant])}-{int(visit_year[participant])}.docx")
        else:
            template.save(f"./Output/feedback_final/{int(adrcid[participant])}-{int(visit_year[participant])}_spanish.docx")
